from docx import Document

from pathlib import Path

import sys



def count_keys(comments, search_terms):
    counts = {key_word: 0 for key_word in search_terms}
    amount_comments = len(comments)

    for comment in comments:
        for key_word in search_terms:
            counts[key_word] += comment.lower().count(key_word.lower())

    print(f"Total comments: {amount_comments}\n")
    for key_word, count in counts.items():
        print(f"{key_word}: {count}\n")

    return amount_comments, counts


    


def number_comments(comments, channel):
    number = input("How many comments do you need?: ")
    while not number.isdigit():
        number = input("\nEnter again: ")

    print(f"\nChannel: {channel}")
    for index, comment in enumerate(comments[:int(number)], 1):
        print(f"\n\n{index}:\n{comment}")







def save_docx(comments, channel, counts, amount_comments, video_id):
    choice = input("\n\nDo you want to save the comments in a DOCX file? (y/n): ").lower()
    while True:
        if choice in ["y", "n"]:
            break
        choice = input("\nEnter again correctly (y/n): ").lower()
    if choice == "n":
        return choice, None

    elif choice == "y":
        if getattr(sys, 'frozen', False):
            exe_path = Path(sys.executable).resolve()
            app_folder = exe_path.parents[1]
        
        else:
            app_folder = Path(__file__).resolve().parents[1]


        youtube_folder = Path(app_folder, "YouTubeComments")
        youtube_folder.mkdir(parents=True, exist_ok=True)

        full_path = Path(youtube_folder, f"{video_id}.docx")

        counter = 0
        while full_path.exists():
            counter += 1
            full_path = Path(youtube_folder, f"{video_id} ({counter}).docx")
            

        doc = Document()
        doc.add_heading("Comment Report for YouTube", 0)
        
        doc.add_paragraph(f"Channel ID: {channel}")
        doc.add_paragraph(f"Total comments found: {amount_comments}")
        
        doc.add_heading("Keyword Statistics:", level=1)
        for key_word, count in counts.items():
            doc.add_paragraph(f"{key_word}: {count}", style='List Bullet')
            
        doc.add_heading("Selected Comments:", level=1)
        for index, comment in enumerate(comments, 1):
            doc.add_paragraph(f"{index}.\n{comment}")
            doc.add_paragraph("-" * 20)
            

        doc.save(full_path)
        print(f'\nFile "{full_path}" saved successfully!')
        
        return choice, full_path