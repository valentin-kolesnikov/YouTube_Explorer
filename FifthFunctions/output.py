from docx import Document

from pathlib import Path

from os import makedirs, path

import sys



def available_languages(transcript_subtitles):
    language = transcript_subtitles.translation_languages

    available_lang = ", ".join(
        f"{languages.language} ({languages.language_code})"
        for languages in language
    )

    return available_lang




def transcript_fetch(transcript_subtitles):
    snippets = transcript_subtitles.fetch()

    full_text = "\n\n".join(snippet.text for snippet in snippets)

    return full_text




def output(transcript_subtitles, available_lang, full_text):
    
    print(f"Video URL: https://www.youtube.com/watch?v={transcript_subtitles.video_id}\n"
        f"\nCurrent: {transcript_subtitles.language}, {transcript_subtitles.language_code}\n"
        f"\nAvailable languages: {available_lang}\n"
        "\nTranscript:\n"
        f"\n\n{full_text}\n\n")
    





def save_docx(transcript_subtitles, available_lang, full_text):
    choice = input("\n\nDo you want to save the full transcript in a DOCX file? (y/n): ").lower()
    while True:
        if choice in ["y", "n"]:
            break
        choice = input("\nEnter again correctly (y/n): ").lower()

    if choice == "y":
        if getattr(sys, 'frozen', False):
            exe_path = Path(sys.executable).resolve()
            app_folder = exe_path.parents[1]
        
        else:
            app_folder = Path(__file__).resolve().parents[1]


        youtube_folder = path.join(app_folder, "YouTubeTranscripts")
        makedirs(youtube_folder, exist_ok=True)

        full_path = path.join(youtube_folder, f"{transcript_subtitles.video_id}.docx")

        counter = 0
        while path.exists(full_path):
            counter += 1
            full_path = path.join(youtube_folder, f"{transcript_subtitles.video_id} ({counter}).docx")
            

        doc = Document()
        doc.add_heading("The Video Transcript", 0)

        doc.add_paragraph(f"Video URL: https://www.youtube.com/watch?v={transcript_subtitles.video_id}")
        doc.add_paragraph(f"Current: {transcript_subtitles.language}, {transcript_subtitles.language_code}")
        doc.add_paragraph(f"Available languages: {available_lang}")

        doc.add_heading("Transcript:\n", 1)

        doc.add_paragraph(full_text)

        doc.save(full_path)

        print(f'\nFile "{full_path}" saved successfully!')
        return